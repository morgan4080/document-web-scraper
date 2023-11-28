import os

from docx.opc.exceptions import PackageNotFoundError
from slugify import slugify
from docx import Document
import http.client
import json
import re
import pdfplumber


def extract_excerpt_from_pdf(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        pdf_text = ""
        for page in pdf.pages:
            pdf_text += page.extract_text()
        # Extract the first 300 characters as the excerpt
        excerpt = pdf_text[:300]
        # Check if the excerpt reaches a partial word and adjust to end at a word boundary
        if len(excerpt) >= 300:
            excerpt = excerpt[:excerpt.rfind(' ')] + '...'
        return excerpt


def extract_excerpt(docx_file):
    doc = Document(docx_file)
    word_count = 0
    excerpt = ""
    for paragraph in doc.paragraphs:
        words = re.findall(r'\w+', paragraph.text)
        word_count += len(words)
        excerpt += paragraph.text + " "
        if word_count >= 300:
            break
    return excerpt.strip()


def extract_text_from_pdf(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        pdf_text = ""
        for page in pdf.pages:
            pdf_text += page.extract_text()
        return pdf_text


def convert_pdf_to_html(pdf_file):
    pdf_text = extract_text_from_pdf(pdf_file)
    # Your HTML conversion logic here based on extracted text
    # You might use similar logic to convert the text into HTML as per your requirements
    # This could involve adding paragraph tags, headers, etc.
    html_content = "<p>" + pdf_text.replace("\n", "</p><p>") + "</p>"
    return html_content


def convert_docx_to_html(docx_file):
    doc = Document(docx_file)
    html_content = ""
    in_list = False  # Flag to track if currently in a list
    for paragraph in doc.paragraphs:
        # Check if the paragraph is in a numbered or bulleted list
        if paragraph.style.name.startswith('List'):
            if not in_list:
                # If not in a list, start a new list
                html_content += "<ul>" if paragraph.style.name.startswith('List Paragraph Bullet') else "<ol>"
                in_list = True

            # Add list item
            html_content += f"<li>{paragraph.text}</li>"
        else:
            # End the list if we were previously in one
            if in_list:
                html_content += "</ul>" if paragraph.style.name.startswith('List Paragraph Bullet') else "</ol>"
                in_list = False

            # Extracting headers
            if paragraph.style.name.startswith('Heading'):
                level = int(paragraph.style.name.split(' ')[-1])
                html_content += f"<h{level}>{paragraph.text}</h{level}>"
            else:
                html_content += f"<p>{paragraph.text}</p>"

    html_content += f"<a href='/order/create' class='place_order'>Order Now</a>"
    return html_content


def convert_folder_to_html(input_folder):
    total_files_added = 0
    total_files_found = 0
    for file_name in os.listdir(input_folder):
        total_files_found += 1  # Increment total files found
        if file_name.endswith(".docx"):
            try:
                file_path = os.path.join(input_folder, file_name)
                html_content = convert_docx_to_html(file_path)
                excerpt = extract_excerpt(file_path)
                if len(file_name.split(".")) > 0:
                    send_html_to_api(file_name.split(".")[0], excerpt, html_content)
                    total_files_added += 1
            except PackageNotFoundError as e:
                print("Error:", e)
            finally:
                print("Transfer Complete: ", file_name)
        elif file_name.endswith(".pdf"):
            try:
                file_path = os.path.join(input_folder, file_name)
                html_content = convert_pdf_to_html(file_path)
                excerpt = extract_excerpt_from_pdf(file_path)
                if len(file_name.split(".")) > 0:
                    send_html_to_api(file_name.split(".")[0], excerpt, html_content)
                    total_files_added += 1
            except Exception as e:
                print("Error:", e)
            finally:
                print("Transfer Complete: ", file_name)

    return total_files_added, total_files_found


def send_html_to_api(title, excerpt, html):
    conn = http.client.HTTPConnection("localhost", 3000)
    slug = slugify(title)

    # Load existing processed slugs or create an empty list
    processed_slugs = []
    file_path = 'processed_slugs.json'
    if os.path.exists(file_path):
        with open(file_path, 'r') as file:
            processed_slugs = json.load(file)

    if slug in processed_slugs:
        print(f"Entry with slug '{slug}' already processed. Skipping...")
        return

    payload = json.dumps({
        "title": title,
        "slug": slug,
        "excerpt": excerpt,
        "description": html
    })
    headers = {
        'Content-Type': 'application/json'
    }
    conn.request("POST", "/api/papers", payload, headers)
    res = conn.getresponse()
    data = res.read().decode('utf-8')

    # Extracting the slug from the response data
    try:
        response_json = json.loads(data)
        data_info = response_json.get('data')
        if data_info and 'slug' in data_info:
            created_slug = data_info['slug']
            # Add the processed slug to the list and save it back to the file
            processed_slugs.append(created_slug)
            with open(file_path, 'w') as file:
                json.dump(processed_slugs, file)
    except json.JSONDecodeError as e:
        print(f"Error decoding JSON response: {e}")
